"""
Parses a .docx into a ParsedDocument. Uses python-docx's real heading
styles (Heading 1, Heading 2, ...) for structure — the most reliable
format we handle, since Word stores heading level as explicit metadata
rather than something we have to guess at.
"""
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.ingestion.loaders.base import (
    HeadingEl, ParagraphEl, TableEl, ParsedDocument,
    is_metadata_table, extract_metadata_from_grid,
)


def _heading_level(paragraph: Paragraph) -> int | None:
    style = paragraph.style.name if paragraph.style else ""
    if style == "Title":
        return 0
    if style.startswith("Heading "):
        try:
            return int(style.replace("Heading ", ""))
        except ValueError:
            return None
    return None


def _table_to_grid(table: Table) -> list[list[str]]:
    return [[cell.text.strip() for cell in row.cells] for row in table.rows]


def parse(path: str) -> ParsedDocument:
    doc = Document(path)
    title = ""
    metadata = {}
    elements = []

    body = doc.element.body
    para_map = {p._p: p for p in doc.paragraphs}
    table_map = {t._tbl: t for t in doc.tables}

    for child in body.iterchildren():
        if child.tag.endswith("}p") and child in para_map:
            p = para_map[child]
            text = p.text.strip()
            if not text:
                continue
            level = _heading_level(p)
            if level == 0:
                title = text
            elif level is not None:
                elements.append(HeadingEl(level=level, text=text))
            else:
                elements.append(ParagraphEl(text=text))
        elif child.tag.endswith("}tbl") and child in table_map:
            grid = _table_to_grid(table_map[child])
            if is_metadata_table(grid) and not metadata:
                metadata = extract_metadata_from_grid(grid)
            else:
                elements.append(TableEl(rows=grid))

    return ParsedDocument(title=title, metadata=metadata, elements=elements)