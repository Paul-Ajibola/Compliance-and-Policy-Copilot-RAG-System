"""
Parses a .docx into an ordered list of structural elements:
  - heading elements (with level 1-9)
  - paragraph elements (body text)
  - table elements (as atomic grids, never split)

Also extracts the metadata table (Document Owner / Department / Effective
Date / Lifecycle Status) if present near the top of the document.
"""
from dataclasses import dataclass, field
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


@dataclass
class HeadingEl:
    level: int
    text: str


@dataclass
class ParagraphEl:
    text: str


@dataclass
class TableEl:
    rows: list  # list[list[str]], first row assumed header


@dataclass
class ParsedDocument:
    title: str
    metadata: dict
    elements: list = field(default_factory=list)


METADATA_KEYS = {
    "document owner": "document_owner",
    "department": "department",
    "effective date": "effective_date",
    "lifecycle status": "lifecycle_status",
    "review cycle": "review_cycle",
}


def _heading_level(paragraph: Paragraph) -> int | None:
    style = paragraph.style.name if paragraph.style else ""
    if style == "Title":
        return 0
    if style.startswith("Heading "):
        try:
            return int(style.replace("Heading ", ""))
        except ValueError:
            return None
    return None


def _table_to_grid(table: Table) -> list[list[str]]:
    grid = []
    for row in table.rows:
        grid.append([cell.text.strip() for cell in row.cells])
    return grid


def _is_metadata_table(grid: list[list[str]]) -> bool:
    if not grid:
        return False
    keys_found = sum(1 for row in grid if row and row[0].strip().lower() in METADATA_KEYS)
    return keys_found >= 2


def parse_docx(path: str) -> ParsedDocument:
    doc = Document(path)
    title = ""
    metadata = {}
    elements = []

    body = doc.element.body
    para_map = {p._p: p for p in doc.paragraphs}
    table_map = {t._tbl: t for t in doc.tables}

    for child in body.iterchildren():
        if child.tag.endswith("}p") and child in para_map:
            p = para_map[child]
            text = p.text.strip()
            if not text:
                continue
            level = _heading_level(p)
            if level == 0:
                title = text
            elif level is not None:
                elements.append(HeadingEl(level=level, text=text))
            else:
                elements.append(ParagraphEl(text=text))
        elif child.tag.endswith("}tbl") and child in table_map:
            t = table_map[child]
            grid = _table_to_grid(t)
            if _is_metadata_table(grid) and not metadata:
                for row in grid:
                    if len(row) >= 2:
                        key = METADATA_KEYS.get(row[0].strip().lower())
                        if key:
                            metadata[key] = row[1].strip()
            else:
                elements.append(TableEl(rows=grid))

    return ParsedDocument(title=title, metadata=metadata, elements=elements)


